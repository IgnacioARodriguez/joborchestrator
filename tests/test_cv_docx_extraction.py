from __future__ import annotations

from io import BytesIO

import pytest
from docx import Document
from docx.enum.section import WD_SECTION

from joborchestrator.intelligence.cv_profile_extractor import (
    CVProfileError,
    extract_text_from_cv,
)


def _docx_bytes(document: Document) -> bytes:
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def test_docx_extraction_preserves_simple_paragraph_order() -> None:
    document = Document()
    document.add_paragraph("Ignacio Rodriguez")
    document.add_paragraph("Backend Engineer")
    document.add_paragraph("Python APIs")

    text = extract_text_from_cv("cv.docx", _docx_bytes(document))

    assert text.splitlines() == ["Ignacio Rodriguez", "Backend Engineer", "Python APIs"]


def test_docx_extraction_includes_header_tables_and_footer_in_document_order() -> None:
    document = Document()
    document.sections[0].header.paragraphs[0].text = "ignacio@example.com"
    document.add_paragraph("Professional Experience")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Backend Developer"
    table.cell(0, 1).text = "2022 - Present"
    document.add_paragraph("Built Python APIs.")
    document.sections[0].footer.paragraphs[0].text = "Málaga, Spain"

    text = extract_text_from_cv("cv.docx", _docx_bytes(document))

    assert text.splitlines() == [
        "ignacio@example.com",
        "Professional Experience",
        "Backend Developer | 2022 - Present",
        "Built Python APIs.",
        "Málaga, Spain",
    ]


def test_docx_extraction_deduplicates_linked_section_headers_and_footers() -> None:
    document = Document()
    document.sections[0].header.paragraphs[0].text = "LinkedIn: example"
    document.sections[0].footer.paragraphs[0].text = "Page footer"
    document.add_paragraph("First page")
    document.add_section(WD_SECTION.NEW_PAGE)
    document.add_paragraph("Second page")

    text = extract_text_from_cv("cv.docx", _docx_bytes(document))

    assert text.count("LinkedIn: example") == 1
    assert text.count("Page footer") == 1
    assert "First page" in text
    assert "Second page" in text


def test_docx_extraction_skips_blank_paragraphs_and_table_cells() -> None:
    document = Document()
    document.add_paragraph("")
    document.add_paragraph("Summary")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = ""
    table.cell(0, 1).text = "Python"
    document.add_paragraph("")

    text = extract_text_from_cv("cv.docx", _docx_bytes(document))

    assert text.splitlines() == ["Summary", "Python"]


def test_docx_extraction_rejects_invalid_package() -> None:
    with pytest.raises(CVProfileError, match="could not be read"):
        extract_text_from_cv("cv.docx", b"not-a-docx")
