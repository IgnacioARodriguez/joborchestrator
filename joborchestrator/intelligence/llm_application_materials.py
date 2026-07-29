from __future__ import annotations

import json
import logging
import os
import re
import unicodedata
from dataclasses import asdict, is_dataclass
from io import BytesIO
from typing import Any

import httpx

from joborchestrator.llm.provider import LLMProviderError, ProviderRegistry
from joborchestrator.prompts import active_prompt_version, load_prompt
from joborchestrator.intelligence.llm_costs import estimate_application_kit_tokens, estimate_cost
from joborchestrator.intelligence.cv_profile_extractor import profile_payload_to_candidate_profile
from joborchestrator.ranking.schemas import CandidateProfile
from joborchestrator.ranking.serialization import result_to_dict
from joborchestrator.storage import persistence as db

DEFAULT_MATERIALS_MODEL = os.getenv("OPENAI_MATERIALS_MODEL") or os.getenv("OPENAI_MODEL") or "gpt-5.4-mini"
DEFAULT_NVIDIA_MATERIALS_MODEL = (
    os.getenv("NVIDIA_MATERIALS_MODEL")
    or os.getenv("NVIDIA_RANKING_MODEL")
    or os.getenv("NVIDIA_MODEL")
    or "nvidia/llama-3.3-nemotron-super-49b-v1"
)
NVIDIA_BASE_URL = os.getenv("NVIDIA_BASE_URL") or "https://integrate.api.nvidia.com/v1"
DEFAULT_NVIDIA_MATERIALS_TIMEOUT_SECONDS = float(os.getenv("NVIDIA_MATERIALS_TIMEOUT_SECONDS", "300"))
DEFAULT_MATERIALS_VALIDATION_RETRIES = int(
    os.getenv("MATERIALS_VALIDATION_RETRIES")
    or os.getenv("OPENAI_MATERIALS_VALIDATION_RETRIES", "3")
)
MAX_MATERIALS_VALIDATION_RETRIES = int(os.getenv("MAX_MATERIALS_VALIDATION_RETRIES", "6"))
RECRUITER_MESSAGE_MAX_CHARS = 320
logger = logging.getLogger(__name__)
ROLE_ATTRIBUTION_TECH_TERMS = [
    "Python",
    "JavaScript",
    "TypeScript",
    "SQL",
    "REST APIs",
    "APIs",
    "AWS",
    "EC2",
    "AWS Lambda",
    "API Gateway",
    "DynamoDB",
    "Terraform",
    "CloudFormation",
    "AWS CDK",
    "Kubernetes",
    "FastAPI",
    "Django",
    "Flask",
    "React",
    "Next.js",
    "MongoDB",
    "Redis",
    "PostgreSQL",
    "MySQL",
    "PHP",
    "Docker",
    "Git",
]
BULLET_PREFIXES = ("-", "*", "•", "▪", "◦", "‣", "·")
EXPERIENCE_DENSITY_CHAR_RATIO = 0.45
EXPERIENCE_DENSITY_BULLET_RULES = [
    {"ratio": 0.50, "floor": 4, "cap": 6},
    {"ratio": 0.35, "floor": 3, "cap": 5},
    {"ratio": 0.25, "floor": 1, "cap": 3},
]
EXPERIENCE_DENSITY_PARSE_FAILURE = (
    "ats_cv_text density validation was not applied because base CV experience roles could not be parsed. "
    "Review the base CV experience headings/date format before trusting compression checks."
)


class LLMMaterialsError(RuntimeError):
    def __init__(
        self,
        message: str,
        *,
        generation_metadata: dict[str, Any] | None = None,
    ) -> None:
        super().__init__(message)
        self.generation_metadata = generation_metadata or {
            "validation_attempts": 0,
            "validation_errors": [],
        }


def estimate_materials_cost(
    job_count: int,
    model: str = DEFAULT_MATERIALS_MODEL,
    *,
    batch: bool = False,
    avg_description_chars: int = 7000,
) -> float:
    input_tokens, output_tokens = estimate_application_kit_tokens(job_count, avg_description_chars)
    return estimate_cost(input_tokens, output_tokens, model, batch=batch)


def materials_prompt_versions() -> dict[str, str]:
    return {
        "materials/nvidia_cv_contract": active_prompt_version("materials", "nvidia_cv_contract"),
        "materials/nvidia_kit_contract": active_prompt_version("materials", "nvidia_kit_contract"),
    }


def build_application_kit_with_llm(
    job: Any,
    ranking: Any | None = None,
    *,
    model: str | None = None,
    api_key: str | None = None,
    timeout: float = 60.0,
    validation_retry_limit: int | None = None,
) -> dict[str, str]:
    key = api_key or os.getenv("OPENAI_API_KEY")
    if not key:
        raise LLMMaterialsError("OPENAI_API_KEY is required to generate materials with API.")

    payload = _materials_payload(job, ranking)
    kit_response = _call_openai(
        payload,
        key,
        model or DEFAULT_MATERIALS_MODEL,
        timeout,
        validation_retry_limit=validation_retry_limit,
    )
    kit = _kit_from_response(kit_response)
    _attach_generation_metadata(kit, kit_response)
    return kit


def build_application_kit_with_nvidia(
    job: Any,
    ranking: Any | None = None,
    *,
    model: str | None = None,
    api_key: str | None = None,
    timeout: float = DEFAULT_NVIDIA_MATERIALS_TIMEOUT_SECONDS,
    validation_retry_limit: int | None = None,
) -> dict[str, str]:
    key = api_key or os.getenv("NVIDIA_API_KEY") or os.getenv("NIM_API_KEY")
    if not key:
        raise LLMMaterialsError("NVIDIA_API_KEY or NIM_API_KEY is required to generate materials with NVIDIA.")

    payload = _materials_payload(job, ranking)
    selected_model = model or DEFAULT_NVIDIA_MATERIALS_MODEL
    cv_response = _call_nvidia_cv(
        payload,
        key,
        selected_model,
        timeout,
        validation_retry_limit=validation_retry_limit,
    )
    try:
        kit_response = _call_nvidia_kit(
            payload,
            key,
            selected_model,
            timeout,
            validation_retry_limit=validation_retry_limit,
        )
    except LLMMaterialsError as exc:
        metadata = _combined_generation_metadata(
            [
                cv_response,
                {"_generation_metadata": exc.generation_metadata},
            ]
        )
        raise LLMMaterialsError(str(exc), generation_metadata=metadata) from exc
    response = {**kit_response, **cv_response}
    response["_generation_metadata"] = _combined_generation_metadata([cv_response, kit_response])
    kit = _kit_from_response(response)
    _attach_generation_metadata(kit, response)
    return kit


def _materials_payload(job: Any, ranking: Any | None = None) -> dict[str, Any]:
    profile_payload = db.get_candidate_profile_payload()
    if not profile_payload:
        raise LLMMaterialsError("No candidate profile configured. Upload a CV in Profile before generating materials.")
    profile = CandidateProfile(**profile_payload_to_candidate_profile(profile_payload))
    base_cv_text = str(profile_payload.get("base_cv_text") or "").strip()
    ranking_payload = _ranking_payload(ranking)
    compact_job = _compact_job(_to_dict(job))
    candidate_profile = asdict(profile)
    return {
        "candidate_profile": candidate_profile,
        "base_cv": {
            "filename": profile_payload.get("base_cv_filename") or "",
            "text": base_cv_text[:24000],
        },
        "job": compact_job,
        "ranking": ranking_payload,
        "ranking_constraints": _materials_ranking_constraints(ranking_payload),
        "application_tone_constraints": _materials_tone_constraints(ranking_payload),
        "experience_claim_constraints": _materials_experience_claim_constraints(base_cv_text),
        "ats_fit_analysis": _build_ats_fit_analysis(compact_job, candidate_profile, base_cv_text, ranking_payload),
        "goal": (
            "Generate truthful, editable application materials and a complete ATS-optimized CV for this specific job. "
            "Optimize for ATS filters and fast application workflow without inventing experience."
        ),
        "rules": [
            "Do not invent employers, degrees, certifications, years of experience, tools or projects.",
            "The ats_cv_text field must be a complete rewritten CV, not notes, and must preserve the candidate's real personal details, experience, education, dates, and achievements from base_cv.",
            "Keep the base CV's overall section structure when possible, but rewrite wording and ordering for ATS fit against this job.",
            "If base_cv is empty, produce the best complete CV draft possible from the candidate profile and mark missing source limitations in risk_flags.",
            "Use job requirements as keywords only when the candidate can truthfully claim or position adjacent experience.",
            "Treat ranking_constraints.avoid_overclaiming_terms as forbidden claim families unless base_cv or candidate_profile explicitly supports the term or specific related technology; do not include unsupported avoid terms or aliases in generated materials.",
            "Use application_tone_constraints to calibrate confidence; risky or skip decisions must be cautious and must not sound like an automatic strong fit.",
            "Use ats_fit_analysis as the ATS keyword map: emphasize supported_keywords, frame adjacent_or_review_keywords cautiously, and never claim avoid_keywords as direct experience.",
            "Recruiter_message must be a short LinkedIn connection note to a recruiter or hiring contact, not a cover letter and not multiple variants.",
            f"Recruiter_message must fit a LinkedIn invite: maximum {RECRUITER_MESSAGE_MAX_CHARS} characters, one paragraph, no formal letter salutation, no cover-letter body.",
            "Recruiter_message should say why this specific role matches the CV and that the candidate would like to send/share the CV.",
            "Output language should match the job posting language unless the user profile clearly indicates otherwise.",
            "ATS CV text should be ready to copy, export to DOCX/PDF, and submit after human review.",
            "Cover letter is required and must be substantive; when constraints are risky, write it cautiously instead of leaving it empty.",
            "Autofill notes should include copy-paste answers for common portal questions and caveats for claims to avoid.",
            "List risk_flags for unsupported claims, adjacency framing, or user facts to double-check.",
            "Return only JSON matching the schema.",
        ],
        "output_shape": {
            "recruiter_message": "short recruiter connection note, ready to paste into LinkedIn invite/InMail/email",
            "cover_letter": "concise tailored cover letter",
            "ats_cv_text": "complete ATS-optimized CV only; no notes or internal instructions",
            "autofill_notes": "structured copy-paste application workflow",
            "risk_flags": ["unsupported or review-needed claims"],
            "keywords_used": ["truthful job keywords included"],
        },
    }


def _materials_ranking_constraints(ranking: dict[str, Any] | None) -> dict[str, Any]:
    if not ranking:
        return {"avoid_overclaiming_terms": [], "avoid_overclaiming_aliases": {}, "keywords_to_emphasize": []}
    avoid_terms = _terms_from_maybe_json(
        ranking.get("cv_keywords_to_avoid_overclaiming")
        or ranking.get("cv_keywords_to_avoid_overclaiming_json")
    )
    return {
        "avoid_overclaiming_terms": avoid_terms,
        "avoid_overclaiming_aliases": {
            term: _avoid_overclaiming_aliases(term)
            for term in avoid_terms
        },
        "keywords_to_emphasize": _terms_from_maybe_json(
            ranking.get("cv_keywords_to_emphasize") or ranking.get("cv_keywords_to_emphasize_json")
        ),
    }


def _materials_tone_constraints(ranking: dict[str, Any] | None) -> dict[str, Any]:
    if not ranking:
        return {
            "ranking_decision": "",
            "risk_terms": [],
            "tone": "standard",
            "forbidden_phrases": [],
            "allowed_phrases": [],
            "rewrite_strategy": "",
        }
    evidence = ranking.get("evidence") if isinstance(ranking.get("evidence"), dict) else {}
    risk_terms = _dedupe_strings(
        [
            *[str(item) for item in evidence.get("dealbreakers") or []],
            *[str(item) for item in evidence.get("red_flags") or []],
            *[str(item) for item in evidence.get("missing_requirements") or []],
        ]
    )
    decision = str(ranking.get("decision") or "").strip().upper()
    cautious = decision in {"SKIP", "AVOID"} or bool(risk_terms)
    return {
        "ranking_decision": decision,
        "risk_terms": risk_terms,
        "tone": "cautious_review" if cautious else "standard",
        "forbidden_phrases": _cautious_tone_forbidden_phrases() if cautious else [],
        "allowed_phrases": _cautious_tone_allowed_phrases() if cautious else [],
        "rewrite_strategy": (
            "exploratory_review: describe supported background as worth reviewing, ask to discuss fit/scope, "
            "and avoid enthusiasm or strong-match sales language."
            if cautious
            else "standard_application"
        ),
    }


def _build_ats_fit_analysis(
    job: dict[str, Any],
    candidate_profile: dict[str, Any],
    base_cv_text: str,
    ranking: dict[str, Any] | None,
) -> dict[str, Any]:
    ranking = ranking or {}
    source_text = _normalize_for_match(
        "\n".join(
            [
                base_cv_text,
                json.dumps(candidate_profile, ensure_ascii=False),
            ]
        )
    )
    job_terms = _ats_candidate_terms(job, ranking)
    avoid_terms = _terms_from_maybe_json(
        ranking.get("cv_keywords_to_avoid_overclaiming")
        or ranking.get("cv_keywords_to_avoid_overclaiming_json")
    )
    avoid_aliases = {
        term: _avoid_overclaiming_aliases(term)
        for term in avoid_terms
    }

    supported: list[str] = []
    adjacent_or_review: list[str] = []
    avoid: list[str] = []
    for term in job_terms:
        if _term_matches_any_avoid_alias(term, avoid_aliases):
            avoid.append(term)
        elif _contains_phrase_for_materials(source_text, term):
            supported.append(term)
        else:
            adjacent_or_review.append(term)

    return {
        "supported_keywords": supported[:30],
        "adjacent_or_review_keywords": adjacent_or_review[:20],
        "avoid_keywords": avoid[:20],
        "rules": [
            "Use supported_keywords in Summary, Technical Skills, or matching Experience bullets when truthful.",
            "Do not list adjacent_or_review_keywords as direct skills; mention only as review-needed context if useful.",
            "Do not include avoid_keywords or their aliases unless directly supported by base_cv/candidate_profile.",
        ],
    }


def _ats_candidate_terms(job: dict[str, Any], ranking: dict[str, Any]) -> list[str]:
    description = "\n".join(
        str(job.get(key) or "")
        for key in ["title", "description_text"]
    )
    normalized_description = _normalize_for_match(description)
    terms: list[str] = []
    terms.extend(
        term
        for term in ROLE_ATTRIBUTION_TECH_TERMS
        if _contains_phrase_for_materials(normalized_description, term)
    )
    terms.extend(
        _terms_from_maybe_json(
            ranking.get("cv_keywords_to_emphasize")
            or ranking.get("cv_keywords_to_emphasize_json")
        )
    )
    terms.extend(
        _terms_from_maybe_json(
            ranking.get("cv_keywords_to_avoid_overclaiming")
            or ranking.get("cv_keywords_to_avoid_overclaiming_json")
        )
    )
    return _dedupe_strings(terms)


def _term_matches_any_avoid_alias(term: str, avoid_aliases: dict[str, list[str]]) -> bool:
    normalized_term = _normalize_for_match(term)
    for aliases in avoid_aliases.values():
        if any(
            _contains_phrase_for_materials(normalized_term, alias)
            or _contains_phrase_for_materials(_normalize_for_match(alias), term)
            for alias in aliases
        ):
            return True
    return False


def _materials_experience_claim_constraints(base_cv_text: str) -> list[dict[str, Any]]:
    entries = _extract_base_experience_entries(base_cv_text)
    section = _experience_section(base_cv_text)
    constraints = []
    for entry in entries:
        block = _experience_block_for_entry(section, entry, entries)
        supported_technologies = [
            term
            for term in ROLE_ATTRIBUTION_TECH_TERMS
            if _contains_phrase_for_materials(block, term)
        ]
        canonical_technologies = _canonical_role_technologies(block)
        constraints.append(
            {
                "employer": entry["company"],
                "title": entry["title"],
                "supported_role_technologies": supported_technologies,
                "canonical_role_technologies": canonical_technologies,
                "rule": "Inside this employer's bullets or Technologies line, use only technologies supported by this employer block; preserve canonical_role_technologies when present.",
            }
        )
    return constraints


def _kit_from_response(response: dict[str, Any]) -> dict[str, str]:
    return {
        "recruiter_message": _clean_recruiter_message(_material_text(response["recruiter_message"])),
        "cover_letter": str(response.get("cover_letter") or ""),
        "ats_cv_text": _clean_cv_text_for_export(str(response["ats_cv_text"])),
        "autofill_notes": _material_text(response["autofill_notes"]),
    }


def _attach_generation_metadata(kit: dict[str, Any], response: dict[str, Any]) -> None:
    metadata = response.get("_generation_metadata")
    if isinstance(metadata, dict):
        kit["_generation_metadata"] = metadata


def _combined_generation_metadata(responses: list[dict[str, Any]]) -> dict[str, Any]:
    attempts = 0
    errors: list[str] = []
    for response in responses:
        metadata = response.get("_generation_metadata")
        if not isinstance(metadata, dict):
            continue
        attempts += int(metadata.get("validation_attempts") or 0)
        errors.extend(str(error) for error in metadata.get("validation_errors") or [])
    return {"validation_attempts": attempts or 1, "validation_errors": errors}


def _materials_validation_retry_limit(payload: dict[str, Any]) -> int:
    retries = DEFAULT_MATERIALS_VALIDATION_RETRIES
    ranking_constraints = payload.get("ranking_constraints") if isinstance(payload.get("ranking_constraints"), dict) else {}
    avoid_terms = ranking_constraints.get("avoid_overclaiming_terms") or []
    if avoid_terms:
        retries += min(2, len(avoid_terms))

    tone = payload.get("application_tone_constraints") if isinstance(payload.get("application_tone_constraints"), dict) else {}
    if tone.get("tone") == "cautious_review":
        retries += 1

    experience_constraints = payload.get("experience_claim_constraints")
    if isinstance(experience_constraints, list) and any(
        constraint.get("canonical_role_technologies")
        for constraint in experience_constraints
        if isinstance(constraint, dict)
    ):
        retries += 1
    return min(retries, MAX_MATERIALS_VALIDATION_RETRIES)


def _coerce_validation_retry_limit(validation_retry_limit: int | None, payload: dict[str, Any]) -> int:
    if validation_retry_limit is None:
        return _materials_validation_retry_limit(payload)
    return max(0, int(validation_retry_limit))


def _validation_failure_metadata(attempt: int, validation_errors: list[str], validation_feedback: str) -> dict[str, Any]:
    metadata = {
        "validation_attempts": attempt + 1,
        "validation_errors": [*validation_errors, validation_feedback],
    }
    if _is_unrecoverable_validation_feedback(validation_feedback):
        metadata["human_review_required"] = True
    return metadata


def _is_unrecoverable_validation_feedback(validation_feedback: str | None) -> bool:
    normalized = _normalize_for_match(validation_feedback or "")
    return _normalize_for_match(EXPERIENCE_DENSITY_PARSE_FAILURE) in normalized


def _request_failure_metadata(attempt: int, validation_errors: list[str], exc: Exception) -> dict[str, Any]:
    return {
        "validation_attempts": attempt + 1,
        "validation_errors": [
            *validation_errors,
            f"request failed during validation attempt {attempt + 1}: {exc}",
        ],
    }


def _clean_recruiter_message(text: str) -> str:
    message = re.sub(r"\n{3,}", "\n\n", str(text or "")).strip()
    if not message:
        return ""
    contamination_patterns = [
        r"\bDear\s+(?:Hiring Manager|Recruiter|Sir/Madam|Team)\b[:,]?",
        r"\bI'?m reaching out to express (?:my )?interest\b",
        r"\bI am writing to (?:express|apply)\b",
        r"\bSincerely\b[:,]?",
        r"\bBest regards\b[:,]?",
    ]
    earliest: int | None = None
    for pattern in contamination_patterns:
        match = re.search(pattern, message, flags=re.IGNORECASE)
        if match and match.start() > 0:
            earliest = match.start() if earliest is None else min(earliest, match.start())
    if earliest is not None:
        message = message[:earliest].strip()
    return re.sub(r"[ \t]+", " ", message).strip()


def _material_text(value: Any) -> str:
    if isinstance(value, dict):
        preferred_keys = ["short", "long", "summary", "copy_paste_block", "notes"]
        parts = [str(value[key]).strip() for key in preferred_keys if str(value.get(key) or "").strip()]
        if parts:
            return "\n\n".join(parts)
        return json.dumps(value, ensure_ascii=False, indent=2)
    if isinstance(value, list):
        return "\n".join(str(item).strip() for item in value if str(item).strip())
    return str(value or "")


def _ranking_payload(ranking: Any | None) -> dict[str, Any] | None:
    if ranking is None:
        return None
    if isinstance(ranking, dict):
        return ranking
    return result_to_dict(ranking)


def build_ats_cv_with_nvidia(
    job: Any,
    ranking: Any | None = None,
    *,
    model: str | None = None,
    api_key: str | None = None,
    timeout: float = DEFAULT_NVIDIA_MATERIALS_TIMEOUT_SECONDS,
) -> dict[str, Any]:
    key = api_key or os.getenv("NVIDIA_API_KEY") or os.getenv("NIM_API_KEY")
    if not key:
        raise LLMMaterialsError("NVIDIA_API_KEY or NIM_API_KEY is required to generate materials with NVIDIA.")
    payload = _materials_payload(job, ranking)
    return _call_nvidia_cv(payload, key, model or DEFAULT_NVIDIA_MATERIALS_MODEL, timeout)


def build_lightweight_kit_with_nvidia(
    job: Any,
    ranking: Any | None = None,
    *,
    model: str | None = None,
    api_key: str | None = None,
    timeout: float = DEFAULT_NVIDIA_MATERIALS_TIMEOUT_SECONDS,
) -> dict[str, Any]:
    key = api_key or os.getenv("NVIDIA_API_KEY") or os.getenv("NIM_API_KEY")
    if not key:
        raise LLMMaterialsError("NVIDIA_API_KEY or NIM_API_KEY is required to generate materials with NVIDIA.")
    payload = _materials_payload(job, ranking)
    return _call_nvidia_kit(payload, key, model or DEFAULT_NVIDIA_MATERIALS_MODEL, timeout)


def export_ats_cv_docx_bytes(job: dict[str, Any], ats_cv_text: str) -> bytes:
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt
    except ModuleNotFoundError as exc:
        raise LLMMaterialsError("DOCX export requires python-docx. Install it with `pip install python-docx`.") from exc

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    styles = document.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(10)

    lines = _clean_cv_text_for_export(ats_cv_text).splitlines()
    for index, block in enumerate(lines):
        text = block.strip()
        if not text:
            document.add_paragraph("")
        elif index == 0:
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(text)
            run.bold = True
            run.font.size = Pt(20)
        elif _is_cv_section_heading(text):
            paragraph = document.add_paragraph()
            run = paragraph.add_run(text.upper())
            run.bold = True
            run.font.size = Pt(12)
            paragraph.paragraph_format.space_before = Pt(8)
            paragraph.paragraph_format.space_after = Pt(2)
        elif text.startswith(("-", "*")):
            paragraph = document.add_paragraph(text[1:].strip(), style="List Bullet")
            paragraph.paragraph_format.left_indent = Inches(0.25)
            paragraph.paragraph_format.space_after = Pt(1)
        elif _looks_like_experience_header(text):
            paragraph = document.add_paragraph()
            run = paragraph.add_run(text)
            run.bold = True
            paragraph.paragraph_format.space_before = Pt(4)
            paragraph.paragraph_format.space_after = Pt(1)
        else:
            paragraph = document.add_paragraph(text)
            if index == 1:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(2)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def export_ats_cv_pdf_bytes(job: dict[str, Any], ats_cv_text: str) -> bytes:
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import cm
        from reportlab.pdfgen import canvas
    except ModuleNotFoundError as exc:
        raise LLMMaterialsError("PDF export requires reportlab. Install it with `pip install reportlab`.") from exc

    buffer = BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4
    left = 2 * cm
    right = width - 2 * cm
    y = height - 1.8 * cm
    lines = _clean_cv_text_for_export(ats_cv_text).splitlines()
    for index, raw_line in enumerate(lines):
        line = raw_line.rstrip()
        if not line:
            y -= 0.22 * cm
            continue
        if index == 0:
            y = _ensure_pdf_space(pdf, y, height, 1.0 * cm)
            pdf.setFont("Times-Bold", 22)
            pdf.drawCentredString(width / 2, y, line)
            y -= 0.55 * cm
            continue
        if index == 1:
            y = _ensure_pdf_space(pdf, y, height, 0.6 * cm)
            pdf.setFont("Times-Roman", 10)
            pdf.drawCentredString(width / 2, y, line)
            y -= 0.55 * cm
            continue
        if _is_cv_section_heading(line):
            y = _ensure_pdf_space(pdf, y, height, 0.8 * cm)
            y -= 0.15 * cm
            pdf.setFont("Times-Bold", 12)
            pdf.drawString(left, y, line.upper())
            pdf.line(left, y - 0.08 * cm, right, y - 0.08 * cm)
            y -= 0.42 * cm
            continue

        bullet = line.lstrip().startswith(("-", "*"))
        x = left + (0.45 * cm if bullet else 0)
        prefix = "- " if bullet else ""
        content = line.lstrip("-* ").strip() if bullet else line
        font = "Times-Bold" if _looks_like_experience_header(line) else "Times-Roman"
        pdf.setFont(font, 10)
        wrapped = _wrap_pdf_line(prefix + content, max_chars=92 if bullet else 98)
        for chunk_index, chunk in enumerate(wrapped):
            y = _ensure_pdf_space(pdf, y, height, 0.45 * cm)
            pdf.setFont(font, 10)
            if bullet and chunk_index > 0:
                chunk = "  " + chunk
            pdf.drawString(x, y, chunk)
            y -= 0.36 * cm
    pdf.save()
    return buffer.getvalue()


def _clean_cv_text_for_export(text: str) -> str:
    cleaned = str(text or "")
    replacements = {
        "\x7f": "-",
        "\u2022": "-",
        "\u2023": "-",
        "\u25e6": "-",
    }
    for old, new in replacements.items():
        cleaned = cleaned.replace(old, new)
    forbidden_sections = [
        "Optimization notes",
        "ATS CV targeting notes",
        "ATS optimized CV draft",
        "Optimized CV",
    ]
    lines = []
    skip_rest = False
    for raw_line in cleaned.splitlines():
        stripped = raw_line.strip()
        if any(stripped.lower().startswith(section.lower()) for section in forbidden_sections):
            if stripped.lower().startswith("optimization notes"):
                skip_rest = True
            continue
        if skip_rest:
            continue
        if stripped.startswith("Target role:") or stripped.startswith("Positioning angle:"):
            continue
        if stripped.startswith("ATS keywords to emphasize truthfully:"):
            continue
        if set(stripped) <= {"-"}:
            continue
        lines.append(raw_line)
    return "\n".join(lines).strip()


def _wrap_pdf_line(line: str, max_chars: int) -> list[str]:
    words = line.split()
    if not words:
        return [""]
    lines: list[str] = []
    current = ""
    for word in words:
        candidate = f"{current} {word}".strip()
        if len(candidate) <= max_chars:
            current = candidate
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def _ensure_pdf_space(pdf: Any, y: float, height: float, needed: float) -> float:
    from reportlab.lib.units import cm

    if y >= 1.8 * cm + needed:
        return y
    pdf.showPage()
    return height - 1.8 * cm


def _is_cv_section_heading(text: str) -> bool:
    normalized = _normalize_for_match(text).strip(" :")
    return normalized in {
        "summary",
        "professional summary",
        "technical skills",
        "skills",
        "experience",
        "professional experience",
        "education",
    }


def _looks_like_experience_header(text: str) -> bool:
    if text.strip().startswith(("-", "*")):
        return False
    normalized = _normalize_for_match(text)
    if "|" in text and any(role in normalized for role in ["developer", "engineer", "consultant", "architect"]):
        return True
    return bool(re.search(r"(?i)\b(developer|engineer|consultant|architect)\b", text) and _date_range_match(text))


def _call_openai(
    payload: dict[str, Any],
    api_key: str,
    model: str,
    timeout: float,
    *,
    validation_retry_limit: int | None = None,
) -> dict[str, Any]:
    validation_feedback: str | None = None
    validation_errors: list[str] = []
    retry_limit = _coerce_validation_retry_limit(validation_retry_limit, payload)
    for attempt in range(retry_limit + 1):
        try:
            parsed = _call_openai_once(payload, api_key, model, timeout, validation_feedback)
        except LLMMaterialsError as exc:
            raise LLMMaterialsError(
                str(exc),
                generation_metadata=_request_failure_metadata(attempt, validation_errors, exc),
            ) from exc
        validation_feedback = _materials_validation_error(parsed, _base_cv_text(payload), payload)
        if not validation_feedback:
            parsed["_generation_metadata"] = {
                "validation_attempts": attempt + 1,
                "validation_errors": validation_errors,
            }
            return parsed
        if attempt < retry_limit and not _is_unrecoverable_validation_feedback(validation_feedback):
            validation_errors.append(validation_feedback)
            logger.warning("Retrying OpenAI materials generation after invalid response: %s", validation_feedback)
            continue
        raise LLMMaterialsError(
            f"OpenAI materials response was incomplete: {validation_feedback}",
            generation_metadata=_validation_failure_metadata(attempt, validation_errors, validation_feedback),
        )
    raise LLMMaterialsError("OpenAI materials response did not produce a usable application kit.")


def _call_nvidia(
    payload: dict[str, Any],
    api_key: str,
    model: str,
    timeout: float,
    *,
    validation_retry_limit: int | None = None,
) -> dict[str, Any]:
    validation_feedback: str | None = None
    validation_errors: list[str] = []
    retry_limit = _coerce_validation_retry_limit(validation_retry_limit, payload)
    for attempt in range(retry_limit + 1):
        try:
            parsed = _call_nvidia_once(payload, api_key, model, timeout, validation_feedback)
        except LLMMaterialsError as exc:
            raise LLMMaterialsError(
                str(exc),
                generation_metadata=_request_failure_metadata(attempt, validation_errors, exc),
            ) from exc
        validation_feedback = _materials_validation_error(parsed, _base_cv_text(payload), payload)
        if not validation_feedback:
            parsed["_generation_metadata"] = {
                "validation_attempts": attempt + 1,
                "validation_errors": validation_errors,
            }
            return parsed
        if attempt < retry_limit and not _is_unrecoverable_validation_feedback(validation_feedback):
            validation_errors.append(validation_feedback)
            logger.warning("Retrying NVIDIA materials generation after invalid response: %s", validation_feedback)
            continue
        raise LLMMaterialsError(
            f"NVIDIA materials response was incomplete: {validation_feedback}",
            generation_metadata=_validation_failure_metadata(attempt, validation_errors, validation_feedback),
        )
    raise LLMMaterialsError("NVIDIA materials response did not produce a usable application kit.")


def _call_nvidia_cv(
    payload: dict[str, Any],
    api_key: str,
    model: str,
    timeout: float,
    *,
    validation_retry_limit: int | None = None,
) -> dict[str, Any]:
    validation_feedback: str | None = None
    validation_errors: list[str] = []
    retry_limit = _coerce_validation_retry_limit(validation_retry_limit, payload)
    for attempt in range(retry_limit + 1):
        try:
            parsed = _call_nvidia_contract_once(
                _nvidia_cv_contract(),
                payload,
                api_key,
                model,
                timeout,
                validation_feedback,
            )
        except LLMMaterialsError as exc:
            raise LLMMaterialsError(
                str(exc),
                generation_metadata=_request_failure_metadata(attempt, validation_errors, exc),
            ) from exc
        validation_feedback = _ats_cv_response_validation_error(parsed, _base_cv_text(payload), payload)
        if not validation_feedback:
            parsed["_generation_metadata"] = {
                "validation_attempts": attempt + 1,
                "validation_errors": validation_errors,
            }
            return parsed
        if attempt < retry_limit and not _is_unrecoverable_validation_feedback(validation_feedback):
            validation_errors.append(validation_feedback)
            logger.warning(
                "Retrying NVIDIA ATS CV generation after invalid response: %s received_keys=%s",
                validation_feedback,
                sorted(parsed.keys()),
            )
            continue
        raise LLMMaterialsError(
            f"NVIDIA ATS CV response was incomplete: {validation_feedback}",
            generation_metadata=_validation_failure_metadata(attempt, validation_errors, validation_feedback),
        )
    raise LLMMaterialsError("NVIDIA ATS CV response did not produce a usable CV.")


def _call_nvidia_kit(
    payload: dict[str, Any],
    api_key: str,
    model: str,
    timeout: float,
    *,
    validation_retry_limit: int | None = None,
) -> dict[str, Any]:
    validation_feedback: str | None = None
    validation_errors: list[str] = []
    retry_limit = _coerce_validation_retry_limit(validation_retry_limit, payload)
    for attempt in range(retry_limit + 1):
        try:
            parsed = _call_nvidia_contract_once(
                _nvidia_kit_contract(),
                payload,
                api_key,
                model,
                timeout,
                validation_feedback,
            )
        except LLMMaterialsError as exc:
            raise LLMMaterialsError(
                str(exc),
                generation_metadata=_request_failure_metadata(attempt, validation_errors, exc),
            ) from exc
        validation_feedback = _kit_validation_error(parsed, payload)
        if not validation_feedback:
            parsed["_generation_metadata"] = {
                "validation_attempts": attempt + 1,
                "validation_errors": validation_errors,
            }
            return parsed
        if attempt < retry_limit and not _is_unrecoverable_validation_feedback(validation_feedback):
            validation_errors.append(validation_feedback)
            logger.warning(
                "Retrying NVIDIA kit generation after invalid response: %s received_keys=%s",
                validation_feedback,
                sorted(parsed.keys()),
            )
            continue
        raise LLMMaterialsError(
            f"NVIDIA kit response was incomplete: {validation_feedback}",
            generation_metadata=_validation_failure_metadata(attempt, validation_errors, validation_feedback),
        )
    raise LLMMaterialsError("NVIDIA kit response did not produce usable materials.")


def _call_nvidia_contract_once(
    contract: str,
    payload: dict[str, Any],
    api_key: str,
    model: str,
    timeout: float,
    validation_feedback: str | None = None,
) -> dict[str, Any]:
    try:
        provider = ProviderRegistry().get(
            "materials",
            provider_name="nvidia",
            api_key=api_key,
            base_url=NVIDIA_BASE_URL,
            timeout=timeout,
            http_module=httpx,
        )
        response = provider.complete(
            _nvidia_contract_messages(contract, payload, validation_feedback),
            model=model,
            temperature=0,
            response_format="json",
            max_tokens=int(os.getenv("NVIDIA_MATERIALS_MAX_TOKENS", "8000")),
            top_p=0.95,
            frequency_penalty=0,
            presence_penalty=0,
        )
    except LLMProviderError as exc:
        raise LLMMaterialsError(f"NVIDIA materials request failed: {exc}") from exc

    try:
        return json.loads(_extract_json_object_text(response.text))
    except json.JSONDecodeError as exc:
        raise LLMMaterialsError(f"NVIDIA materials response was not valid JSON: {exc}") from exc


def _call_nvidia_once(
    payload: dict[str, Any],
    api_key: str,
    model: str,
    timeout: float,
    validation_feedback: str | None = None,
) -> dict[str, Any]:
    user_payload = dict(payload)
    if validation_feedback:
        user_payload["previous_response_error"] = validation_feedback
        user_payload["instruction"] = "Return a corrected complete JSON object only."
    try:
        provider = ProviderRegistry().get(
            "materials",
            provider_name="nvidia",
            api_key=api_key,
            base_url=NVIDIA_BASE_URL,
            timeout=timeout,
            http_module=httpx,
        )
        response = provider.complete(
            _nvidia_materials_messages(user_payload),
            model=model,
            temperature=0.1,
            response_format="json",
            max_tokens=int(os.getenv("NVIDIA_MATERIALS_MAX_TOKENS", "12000")),
            top_p=0.95,
        )
    except LLMProviderError as exc:
        raise LLMMaterialsError(f"NVIDIA materials request failed: {exc}") from exc

    try:
        return json.loads(_extract_json_object_text(response.text))
    except json.JSONDecodeError as exc:
        raise LLMMaterialsError(f"NVIDIA materials response was not valid JSON: {exc}") from exc


def _call_openai_once(
    payload: dict[str, Any],
    api_key: str,
    model: str,
    timeout: float,
    validation_feedback: str | None = None,
) -> dict[str, Any]:
    user_payload = dict(payload)
    if validation_feedback:
        user_payload["previous_response_error"] = validation_feedback
        user_payload["instruction"] = "Return a corrected complete JSON object only."
    try:
        provider = ProviderRegistry().get(
            "materials",
            provider_name="openai",
            api_key=api_key,
            timeout=timeout,
            http_module=httpx,
        )
        response = provider.complete(
            _openai_materials_messages(user_payload),
            model=model,
            response_format="json",
            response_schema=_materials_schema(),
            schema_name="application_kit",
        )
    except LLMProviderError as exc:
        raise LLMMaterialsError(f"OpenAI materials request failed: {exc}") from exc

    try:
        return json.loads(response.text)
    except json.JSONDecodeError as exc:
        raise LLMMaterialsError("OpenAI materials response was not valid JSON.") from exc


def _nvidia_contract_messages(
    contract: str,
    payload: dict[str, Any],
    validation_feedback: str | None = None,
) -> list[dict[str, Any]]:
    user_content = contract + "\n\nContext:\n" + json.dumps(payload, ensure_ascii=False)
    if validation_feedback:
        repair_instruction = _materials_repair_instruction(validation_feedback)
        user_content += (
            "\n\nYour previous response was rejected because: "
            f"{validation_feedback}\n{repair_instruction}\nReturn a corrected complete JSON object only."
        )
    return [
        {
            "role": "system",
            "content": (
                "You are a strict career application assistant. Return only JSON that matches "
                "the requested shape. Do not include markdown fences or commentary."
            ),
        },
        {"role": "user", "content": user_content},
    ]


def _nvidia_materials_messages(user_payload: dict[str, Any]) -> list[dict[str, Any]]:
    return [
        {
            "role": "system",
            "content": (
                "You are a strict career application assistant. Return only valid JSON. "
                "The ats_cv_text value must be a final complete CV, not notes, comments, or instructions."
            ),
        },
        {"role": "user", "content": json.dumps(user_payload, ensure_ascii=False)},
    ]


def _materials_repair_instruction(validation_feedback: str) -> str:
    normalized = _normalize_for_match(validation_feedback)
    instructions = ["Fix only the rejected fields while preserving all required JSON keys."]
    if "overconfident tone" in normalized:
        instructions.append(
            "For cautious/review rankings, fully rewrite recruiter_message, cover_letter, and autofill_notes in "
            "exploratory-review mode. Use only neutral phrases such as 'may be relevant to review', "
            "'supported Python/API background', 'worth discussing if the contract context fits', and "
            "'I would treat this as exploratory'. Remove the exact words confident, excited, eager, strong fit, "
            "ideal fit, perfect fit, excellent fit, and immediate impact."
        )
    if "ats_cv_text is too short" in normalized or "too few parseable lines" in normalized:
        instructions.append(
            "Rewrite ats_cv_text as a complete ATS CV, not a summary. Include contact/header, Professional Summary, "
            "Technical Skills, Professional Experience with every base CV employer, and Education. Use at least "
            "700 characters and 18 non-empty lines while preserving only truthful source-backed facts."
        )
    if "overcompressed" in normalized:
        instructions.append(
            "Expand ats_cv_text using more source-backed detail from Context.base_cv. Preserve proportionally more "
            "truthful bullets for recent and substantial roles, especially the current and second-most-recent roles. "
            "Do not add unsupported new claims just to increase length."
        )
    if "hedge language" in normalized:
        instructions.append(
            "Remove parenthetical hedges such as 'SQL expertise' or 'implied'. Use broader supported terms "
            "like SQL/databases, or omit the unsupported specific keyword."
        )
    if "internal review/evaluation language" in normalized:
        instructions.append(
            "Remove internal evaluator terms such as ranking, dealbreaker, safety gate, system, validation, "
            "or avoid-overclaiming. Use applicant-facing caveats only."
        )
    if "missing canonical role technologies" in normalized:
        instructions.append(
            "For each employer named in the error, restore the canonical_role_technologies from "
            "Context.experience_claim_constraints in that employer's Professional Experience block, preferably in "
            "a stable Technologies line."
        )
    if "avoid-overclaiming terms" in normalized:
        instructions.append(
            "Remove every forbidden alias from every field, including caveats; describe gaps generically."
        )
    return " ".join(instructions)


def _openai_materials_messages(user_payload: dict[str, Any]) -> list[dict[str, Any]]:
    user_content = _openai_materials_contract() + "\n\nContext:\n" + json.dumps(user_payload, ensure_ascii=False)
    return [
        {
            "role": "system",
            "content": (
                "You are a strict career application assistant. Create high-quality, truthful, ATS-aware "
                "materials. Return only structured JSON. Do not leave required sections blank."
            ),
        },
        {"role": "user", "content": user_content},
    ]


def _extract_json_object_text(text: str) -> str:
    cleaned = text.strip()
    if cleaned.startswith("```"):
        cleaned = cleaned.removeprefix("```json").removeprefix("```").strip()
        cleaned = cleaned.removesuffix("```").strip()
    start = cleaned.find("{")
    end = cleaned.rfind("}")
    if start < 0 or end <= start:
        raise json.JSONDecodeError("No JSON object found", cleaned, 0)
    return cleaned[start : end + 1]


def _materials_schema() -> dict[str, Any]:
    return {
        "type": "object",
        "additionalProperties": False,
        "required": [
            "recruiter_message",
            "cover_letter",
            "ats_cv_text",
            "autofill_notes",
            "risk_flags",
            "keywords_used",
        ],
        "properties": {
            "recruiter_message": {"type": "string"},
            "cover_letter": {"type": "string"},
            "ats_cv_text": {"type": "string"},
            "autofill_notes": {"type": "string"},
            "risk_flags": {"type": "array", "items": {"type": "string"}},
            "keywords_used": {"type": "array", "items": {"type": "string"}},
        },
    }


def _nvidia_cv_contract() -> str:
    return load_prompt("materials", "nvidia_cv_contract")


def _nvidia_kit_contract() -> str:
    return load_prompt("materials", "nvidia_kit_contract")


def _openai_materials_contract() -> str:
    return (
        "ATS CV contract:\n"
        + _nvidia_cv_contract()
        + "\n\nApplication kit contract:\n"
        + _nvidia_kit_contract()
        + "\n\nReturn one JSON object containing the ATS CV fields and application kit fields."
    )


def _materials_validation_error(
    payload: dict[str, Any],
    base_cv_text: str | None = None,
    source_payload: dict[str, Any] | None = None,
) -> str | None:
    if base_cv_text is None and source_payload is not None:
        base_cv_text = _base_cv_text(source_payload)
    problems = []
    kit_error = _kit_response_validation_error(payload, source_payload)
    cv_error = _ats_cv_response_validation_error(payload, base_cv_text, source_payload)
    if kit_error:
        problems.append(kit_error)
    if cv_error:
        problems.append(cv_error)
    non_cv_error = _materials_non_cv_overclaiming_error(payload, source_payload)
    if non_cv_error:
        problems.append(non_cv_error)
    return "; ".join(problems) if problems else None


def _kit_response_validation_error(
    payload: dict[str, Any],
    source_payload: dict[str, Any] | None = None,
) -> str | None:
    problems = []
    for field in ["recruiter_message", "cover_letter", "autofill_notes"]:
        if not str(payload.get(field) or "").strip():
            problems.append(f"{field} is required")
    recruiter_message = str(payload.get("recruiter_message") or "")
    cover_letter = str(payload.get("cover_letter") or "").strip()
    if len(recruiter_message) > RECRUITER_MESSAGE_MAX_CHARS:
        problems.append("recruiter_message is too long")
    if cover_letter and len(cover_letter) < 120:
        problems.append("cover_letter is too short to be substantive")
    problems.extend(_recruiter_message_quality_problems(recruiter_message))
    problems.extend(_recruiter_message_specificity_problems(recruiter_message, source_payload))
    kit_text = "\n".join(str(payload.get(field) or "") for field in ["recruiter_message", "cover_letter", "autofill_notes"])
    problems.extend(_unsupported_hedge_problems(kit_text))
    problems.extend(_materials_internal_note_problems(kit_text))
    problems.extend(_application_tone_problems(payload, source_payload))
    problems.extend(_unsupported_experience_years_problems(kit_text, source_payload, field_name="application_materials"))
    return "; ".join(problems) if problems else None


def _recruiter_message_quality_problems(text: str) -> list[str]:
    message = str(text or "").strip()
    lower = message.lower()
    problems: list[str] = []
    cover_letter_markers = [
        "dear hiring manager",
        "dear recruiter",
        "i am writing to express",
        "i'm writing to express",
        "i am reaching out to express interest",
        "i'm reaching out to express interest",
        "sincerely",
    ]
    found = [marker for marker in cover_letter_markers if marker in lower]
    if found:
        problems.append(f"recruiter_message reads like a cover letter: {', '.join(found[:2])}")
    intro_markers = len(re.findall(r"\b(?:i am|i'm)\s+[^.\n]{0,90}\b(?:developer|engineer|specialist|manager|consultant)\b", lower))
    if intro_markers > 1:
        problems.append("recruiter_message repeats the candidate introduction")
    interest_markers = len(re.findall(r"\b(?:interested in|interest in|excited about|express interest)\b", lower))
    if interest_markers > 2:
        problems.append("recruiter_message repeats the interest statement")
    return problems


def _recruiter_message_specificity_problems(
    text: str,
    source_payload: dict[str, Any] | None,
) -> list[str]:
    terms = _recruiter_specificity_terms(source_payload)
    if not terms:
        return []
    normalized = _normalize_for_match(text)
    if any(term in normalized for term in terms):
        return []
    return ["recruiter_message is generic; mention the target company or role"]


def _recruiter_specificity_terms(source_payload: dict[str, Any] | None) -> list[str]:
    if not source_payload:
        return []
    job = source_payload.get("job") if isinstance(source_payload.get("job"), dict) else {}
    company = _normalize_for_match(str(job.get("company") or ""))
    title = _normalize_for_match(str(job.get("title") or ""))
    terms = []
    if company and company not in {"confidential", "unknown", "none"}:
        terms.append(company)
    title = re.sub(r"\([^)]*\)", " ", title)
    title = re.sub(r"\s+", " ", title).strip()
    if title:
        terms.append(title)
        role_words = [
            word
            for word in title.split()
            if len(word) >= 3 and word not in {"remote", "hybrid", "onsite", "senior", "junior", "lead"}
        ]
        if len(role_words) >= 2:
            terms.append(" ".join(role_words[:3]))
    return _dedupe_strings(terms)


def _dedupe_strings(values: list[str]) -> list[str]:
    seen = set()
    deduped = []
    for value in values:
        key = str(value or "").strip()
        if not key or key in seen:
            continue
        seen.add(key)
        deduped.append(key)
    return deduped


def _ats_cv_response_validation_error(
    payload: dict[str, Any],
    base_cv_text: str | None = None,
    source_payload: dict[str, Any] | None = None,
) -> str | None:
    problems = []
    ats_cv_text = str(payload.get("ats_cv_text") or "")
    if not ats_cv_text.strip():
        problems.append("ats_cv_text is required")
    for field in ["risk_flags", "keywords_used"]:
        if not isinstance(payload.get(field), list):
            problems.append(f"{field} must be an array")
    problems.extend(_ats_cv_quality_problems(ats_cv_text))
    problems.extend(_experience_coverage_problems(str(base_cv_text or ""), ats_cv_text))
    problems.extend(_experience_density_problems(str(base_cv_text or ""), ats_cv_text))
    problems.extend(_experience_technology_attribution_problems(str(base_cv_text or ""), ats_cv_text))
    problems.extend(_avoid_overclaiming_problems(ats_cv_text, source_payload, field_name="ats_cv_text"))
    problems.extend(_unsupported_experience_years_problems(ats_cv_text, source_payload, field_name="ats_cv_text"))
    return "; ".join(problems) if problems else None


def _kit_validation_error(
    payload: dict[str, Any],
    source_payload: dict[str, Any] | None = None,
) -> str | None:
    problems = []
    kit_error = _kit_response_validation_error(payload, source_payload)
    overclaiming_error = _materials_non_cv_overclaiming_error(payload, source_payload)
    if kit_error:
        problems.append(kit_error)
    if overclaiming_error:
        problems.append(overclaiming_error)
    return "; ".join(problems) if problems else None


def _base_cv_text(payload: dict[str, Any]) -> str:
    base_cv = payload.get("base_cv")
    if isinstance(base_cv, dict):
        return str(base_cv.get("text") or "")
    return ""


def _ats_cv_quality_problems(text: str) -> list[str]:
    cleaned = _clean_cv_text_for_export(text)
    normalized = cleaned.lower()
    raw_normalized = str(text or "").lower()
    problems: list[str] = []
    if len(cleaned) < 700:
        problems.append("ats_cv_text is too short to be a complete ATS CV")
    line_count = len([line for line in cleaned.splitlines() if line.strip()])
    if line_count < 18:
        problems.append(f"ats_cv_text has too few parseable lines for a complete CV: {line_count}/18")

    section_patterns = {
        "summary": ["summary", "profile", "professional summary", "perfil", "resumen"],
        "experience": ["experience", "work experience", "professional experience", "experiencia"],
        "skills": ["skills", "technical skills", "core skills", "competencias", "habilidades"],
        "education": ["education", "formacion", "formación", "academic", "educacion", "educación"],
    }
    missing_sections = [
        section
        for section, aliases in section_patterns.items()
        if not any(_contains_section_heading(normalized, alias) for alias in aliases)
    ]
    if missing_sections:
        problems.append(f"ats_cv_text is missing standard ATS sections: {', '.join(missing_sections)}")

    forbidden_markers = [
        "optimization notes",
        "ats cv targeting notes",
        "target role:",
        "positioning angle:",
        "do not add skills",
        "profile-backed keywords",
        "keywords to emphasize",
        "internal note",
    ]
    found_markers = [marker for marker in forbidden_markers if marker in raw_normalized]
    if found_markers:
        problems.append(f"ats_cv_text contains internal/non-CV notes: {', '.join(found_markers[:3])}")
    problems.extend(_unsupported_hedge_problems(text))
    return problems


def _contains_section_heading(normalized_text: str, heading: str) -> bool:
    for line in normalized_text.splitlines():
        stripped = line.strip(" :-\t")
        if stripped == heading or stripped.startswith(f"{heading}:"):
            return True
    return False


def _experience_coverage_problems(base_cv_text: str, ats_cv_text: str) -> list[str]:
    entries = _extract_base_experience_entries(base_cv_text)
    if not entries:
        return []
    normalized_cv = _normalize_for_match(ats_cv_text)
    missing = []
    for entry in entries:
        terms = entry["terms"]
        if not any(term in normalized_cv for term in terms):
            missing.append(entry["label"])
    if missing:
        return [f"ats_cv_text omitted base CV experience entries: {', '.join(missing[:6])}"]
    return []


def _experience_density_problems(base_cv_text: str, ats_cv_text: str) -> list[str]:
    entries = _extract_base_experience_entries(base_cv_text)
    source_section = _experience_section(base_cv_text)
    if not entries:
        if _looks_like_unparsed_experience_text(base_cv_text, source_section):
            return [EXPERIENCE_DENSITY_PARSE_FAILURE]
        return []
    generated_section = _experience_section(ats_cv_text) or ats_cv_text
    if not source_section or not generated_section:
        return [EXPERIENCE_DENSITY_PARSE_FAILURE]

    problems: list[str] = []
    base_chars = len(_normalize_whitespace_for_materials(source_section))
    generated_chars = len(_normalize_whitespace_for_materials(generated_section))
    if base_chars >= 1400 and generated_chars / max(base_chars, 1) < EXPERIENCE_DENSITY_CHAR_RATIO:
        problems.append(
            f"ats_cv_text is overcompressed compared with base CV experience detail: "
            f"{generated_chars}/{base_chars} chars. Preserve more source-backed role detail."
        )

    compressed_roles = []
    for index, entry in enumerate(entries):
        source_block = _experience_block_for_entry(source_section, entry, entries)
        generated_block = _experience_block_for_entry(generated_section, entry, entries)
        if not source_block:
            continue
        if not generated_block:
            compressed_roles.append(f"{entry['company']} is missing from generated experience")
            continue
        source_bullets = _cv_bullet_count(source_block)
        generated_bullets = _cv_bullet_count(generated_block)
        if source_bullets < 1:
            continue
        required_bullets = _minimum_bullets_for_role(index, source_bullets)
        if generated_bullets < required_bullets:
            compressed_roles.append(
                f"{entry['company']} kept {generated_bullets}/{source_bullets} bullets; expected at least {required_bullets}"
            )
    if compressed_roles:
        problems.append(
            "ats_cv_text is overcompressed for base CV experience roles: "
            + "; ".join(compressed_roles[:4])
            + ". Preserve proportionally more truthful bullets for recent or substantial roles."
        )
    return problems[:2]


def _minimum_bullets_for_role(index: int, source_bullets: int) -> int:
    import math

    rule = EXPERIENCE_DENSITY_BULLET_RULES[min(index, len(EXPERIENCE_DENSITY_BULLET_RULES) - 1)]
    proportional = math.ceil(source_bullets * float(rule["ratio"]))
    floor = min(int(rule["floor"]), source_bullets)
    return min(source_bullets, int(rule["cap"]), max(floor, proportional))


def _cv_bullet_count(block: str) -> int:
    return sum(1 for line in str(block or "").splitlines() if line.strip().startswith(BULLET_PREFIXES))


def _looks_like_unparsed_experience_text(base_cv_text: str, source_section: str) -> bool:
    candidate = source_section or base_cv_text
    normalized = _normalize_whitespace_for_materials(candidate)
    if source_section:
        return len(normalized) >= 1400
    date_ranges = len([line for line in str(candidate or "").splitlines() if _date_range_match(line)])
    bullets = _cv_bullet_count(candidate)
    role_terms = len(
        re.findall(
            r"(?i)\b(developer|engineer|consultant|architect|manager|analyst|specialist|lead|desarrollador|ingeniero)\b",
            candidate,
        )
    )
    if date_ranges >= 1 and bullets >= 3:
        return True
    if len(normalized) >= 1400 and date_ranges >= 1 and role_terms >= 2:
        return True
    return False


def _normalize_whitespace_for_materials(text: str) -> str:
    return re.sub(r"\s+", " ", str(text or "")).strip()


def _experience_technology_attribution_problems(base_cv_text: str, ats_cv_text: str) -> list[str]:
    entries = _extract_base_experience_entries(base_cv_text)
    if len(entries) < 2:
        return []
    source_section = _experience_section(base_cv_text)
    generated_section = _experience_section(ats_cv_text) or ats_cv_text
    problems: list[str] = []
    for entry in entries:
        source_block = _experience_block_for_entry(source_section, entry, entries)
        generated_block = _experience_block_for_entry(generated_section, entry, entries)
        if not source_block or not generated_block:
            continue
        unsupported = [
            term
            for term in ROLE_ATTRIBUTION_TECH_TERMS
            if _contains_phrase_for_materials(generated_block, term)
            and not _contains_phrase_for_materials(source_block, term)
        ]
        canonical_terms = _canonical_role_technologies(source_block)
        missing_canonical = [
            term
            for term in canonical_terms
            if not _contains_phrase_for_materials(generated_block, term)
        ]
        if unsupported:
            problems.append(
                f"{entry['company']} has unsupported role-specific technologies: {', '.join(unsupported[:6])}. "
                "Remove those technologies from that employer block; if they are globally supported, keep them only "
                "in Professional Summary or Technical Skills."
            )
        if missing_canonical:
            problems.append(
                f"{entry['company']} is missing canonical role technologies: {', '.join(missing_canonical[:6])}. "
                "Preserve the employer's canonical Technologies set from the base CV to keep historical stack "
                "attribution stable across target jobs."
            )
    return problems[:6]


def _canonical_role_technologies(normalized_or_raw_block: str) -> list[str]:
    tech_lines = [
        line
        for line in str(normalized_or_raw_block or "").splitlines()
        if "technolog" in _normalize_for_match(line)
    ]
    if not tech_lines:
        return []
    text = "\n".join(tech_lines)
    return [
        term
        for term in ROLE_ATTRIBUTION_TECH_TERMS
        if _contains_phrase_for_materials(text, term)
    ]


def _unsupported_hedge_problems(text: str) -> list[str]:
    normalized = _normalize_for_match(text)
    markers = [
        "implied through experience",
        "implied through",
        "implied by",
        "can be implied",
        "adaptability can be implied",
        "sql expertise",
        "react proficiency",
    ]
    found = [marker for marker in markers if marker in normalized]
    if not found:
        return []
    return [
        "generated materials contain ATS-opaque unsupported hedge language: "
        + ", ".join(found[:4])
        + ". State only directly supported skills or use a plain risk flag/caveat."
    ]


def _materials_internal_note_problems(text: str) -> list[str]:
    normalized = _normalize_for_match(text)
    markers = [
        "safety gate",
        "highlighted in your system",
        "system concern",
        "ranking decision",
        "ranking says",
        "dealbreaker",
        "avoid-overclaiming",
        "validation error",
    ]
    found = [marker for marker in markers if marker in normalized]
    if not found:
        return []
    return [
        "application materials expose internal review/evaluation language: "
        + ", ".join(found[:4])
    ]


def _application_tone_problems(
    payload: dict[str, Any],
    source_payload: dict[str, Any] | None,
) -> list[str]:
    if not source_payload:
        return []
    tone = source_payload.get("application_tone_constraints")
    if not isinstance(tone, dict):
        tone = _materials_tone_constraints(
            source_payload.get("ranking") if isinstance(source_payload.get("ranking"), dict) else None
        )
    if tone.get("tone") != "cautious_review":
        return []
    text = _normalize_for_match(
        "\n".join(str(payload.get(field) or "") for field in ["recruiter_message", "cover_letter", "autofill_notes"])
    )
    found = [phrase for phrase in _cautious_tone_forbidden_phrases() if phrase in text]
    if not found:
        return []
    decision = tone.get("ranking_decision") or "risky"
    return [
        f"application materials use overconfident tone for {decision} ranking: "
        + ", ".join(found[:4])
    ]


def _cautious_tone_forbidden_phrases() -> list[str]:
    return [
        "confident my skills",
        "i am confident",
        "immediate impact",
        "strong fit",
        "ideal fit",
        "perfect fit",
        "excited about",
        "excited to",
        "excited to enhance",
        "eager to",
        "eager to enhance",
        "eager to contribute",
        "highly confident",
        "excellent fit",
    ]


def _cautious_tone_allowed_phrases() -> list[str]:
    return [
        "may be relevant to review",
        "supported Python/API background",
        "worth discussing if the contract context fits",
        "I would treat this as exploratory",
        "review whether the scope aligns",
    ]


def _experience_block_for_entry(section_text: str, entry: dict[str, Any], entries: list[dict[str, Any]]) -> str:
    normalized_section = _normalize_for_match(section_text)
    current = _first_entry_position(normalized_section, entry)
    if current is None:
        return ""
    next_positions = [
        position
        for other in entries
        if other is not entry
        for position in [_first_entry_position(normalized_section, other)]
        if position is not None and position > current
    ]
    end = min(next_positions) if next_positions else len(normalized_section)
    return normalized_section[current:end]


def _first_entry_position(normalized_section: str, entry: dict[str, Any]) -> int | None:
    positions = [
        normalized_section.find(term)
        for term in entry.get("terms") or []
        if term and normalized_section.find(term) >= 0
    ]
    return min(positions) if positions else None


def _materials_non_cv_overclaiming_error(
    payload: dict[str, Any],
    source_payload: dict[str, Any] | None,
) -> str | None:
    text = "\n".join(
        str(payload.get(field) or "")
        for field in ["recruiter_message", "cover_letter", "autofill_notes"]
    )
    problems = _avoid_overclaiming_problems(text, source_payload, field_name="application_materials")
    return "; ".join(problems) if problems else None


def _avoid_overclaiming_problems(
    text: str,
    source_payload: dict[str, Any] | None,
    *,
    field_name: str,
) -> list[str]:
    if not source_payload:
        return []
    ranking = source_payload.get("ranking") if isinstance(source_payload.get("ranking"), dict) else {}
    avoid_terms = _terms_from_maybe_json(
        ranking.get("cv_keywords_to_avoid_overclaiming")
        or ranking.get("cv_keywords_to_avoid_overclaiming_json")
        or source_payload.get("cv_keywords_to_avoid_overclaiming")
    )
    if not avoid_terms:
        return []

    normalized_cv = _normalize_for_match(text)
    supported_source = _normalize_for_match(_supported_materials_source_text(source_payload))
    unsupported_terms = []
    for term in avoid_terms:
        matched_aliases = [
            alias
            for alias in _avoid_overclaiming_aliases(term)
            if _contains_phrase_for_materials(normalized_cv, alias)
            and not _contains_phrase_for_materials(supported_source, alias)
        ]
        if matched_aliases:
            unsupported_terms.append(f"{term} ({', '.join(matched_aliases[:4])})")
    if not unsupported_terms:
        return []
    return [
        f"{field_name} contains unsupported ranking avoid-overclaiming terms: "
        + ", ".join(unsupported_terms[:6])
        + ". Remove these exact terms and aliases from every generated field, including caveats or gap notes; "
        + "describe unsupported target-stack gaps generically instead."
    ]


def _unsupported_experience_years_problems(
    text: str,
    source_payload: dict[str, Any] | None,
    *,
    field_name: str,
) -> list[str]:
    if not source_payload:
        return []
    claims = _experience_year_claims(text)
    if not claims:
        return []
    supported_source = _supported_materials_source_text(source_payload)
    supported_normalized = _normalize_for_match(supported_source)
    profile = source_payload.get("candidate_profile") if isinstance(source_payload.get("candidate_profile"), dict) else {}
    real_years = _float_or_none(profile.get("real_experience_years"))
    unsupported: list[str] = []
    for claim in claims:
        if _contains_phrase_for_materials(supported_normalized, claim):
            continue
        claim_years = _experience_year_value(claim)
        if claim_years is not None and real_years is not None and claim_years <= real_years:
            continue
        unsupported.append(claim)
    if not unsupported:
        return []
    return [
        f"{field_name} contains unsupported years-of-experience claims: "
        + ", ".join(unsupported[:4])
        + ". Remove or lower years claims unless Context.base_cv or candidate_profile.real_experience_years supports them."
    ]


def _experience_year_claims(text: str) -> list[str]:
    return _dedupe_strings(
        match.group(0)
        for match in re.finditer(r"\b\d{1,2}\+?\s*(?:years|years'|anos|años)\b", str(text or ""), flags=re.IGNORECASE)
    )


def _experience_year_value(claim: str) -> float | None:
    match = re.search(r"\b(\d{1,2})", str(claim or ""))
    return float(match.group(1)) if match else None


def _float_or_none(value: Any) -> float | None:
    try:
        return float(value)
    except (TypeError, ValueError):
        return None


def _avoid_overclaiming_aliases(term: str) -> list[str]:
    normalized = _normalize_for_match(term)
    aliases = [str(term or "").strip()]
    aliases.extend(
        part.strip()
        for part in re.split(r"[/,;|]", str(term or ""))
        if part.strip()
    )
    if "serverless" in normalized:
        aliases.extend(
            [
                "Serverless",
                "Serverless Architecture",
                "AWS Lambda",
                "Lambda",
                "DynamoDB",
                "API Gateway",
                "EventBridge",
                "SQS",
                "SNS",
                "Step Functions",
                "CloudFormation",
                "AWS CDK",
                "CDK",
            ]
        )
    if any(token in normalized for token in ["terraform", "cloudformation", "cdk"]):
        aliases.extend(["Terraform", "CloudFormation", "AWS CDK", "CDK"])
    if "kubernetes" in normalized:
        aliases.extend(["Kubernetes", "K8s", "EKS", "AKS", "GKE"])
    return _dedupe_strings([alias for alias in aliases if str(alias or "").strip()])


def _supported_materials_source_text(source_payload: dict[str, Any]) -> str:
    base_cv = source_payload.get("base_cv")
    profile = source_payload.get("candidate_profile")
    return "\n".join(
        [
            str(base_cv.get("text") or "") if isinstance(base_cv, dict) else "",
            json.dumps(profile, ensure_ascii=False) if isinstance(profile, dict) else str(profile or ""),
        ]
    )


def _terms_from_maybe_json(value: Any) -> list[str]:
    if isinstance(value, str):
        stripped = value.strip()
        if stripped.startswith("["):
            try:
                return _dedupe_strings([str(item).strip() for item in json.loads(stripped) if str(item).strip()])
            except json.JSONDecodeError:
                pass
        return [stripped] if stripped else []
    if isinstance(value, list):
        return _dedupe_strings([str(item).strip() for item in value if str(item).strip()])
    return []


def _extract_base_experience_entries(base_cv_text: str) -> list[dict[str, Any]]:
    section = _experience_section(base_cv_text)
    if not section:
        return []
    lines = [line.strip() for line in section.splitlines() if line.strip()]
    entries: list[dict[str, Any]] = []
    for index, line in enumerate(lines):
        match = _date_range_match(line)
        if not match:
            continue
        title = line[: match.start()].strip(" -|")
        company = _next_company_line(lines, index + 1)
        if not title or not company:
            continue
        terms = _company_match_terms(company)
        if not terms:
            continue
        entries.append(
            {
                "title": title,
                "company": company,
                "label": f"{title} at {company}",
                "terms": terms,
            }
        )
    return entries


def _experience_section(text: str) -> str:
    headings = (
        "experience",
        "professional experience",
        "work experience",
        "employment history",
        "employment experience",
        "career history",
        "professional background",
        "relevant experience",
        "experiencia",
        "experiencia profesional",
        "historial laboral",
        "trayectoria profesional",
    )
    stop_headings = (
        "projects?",
        "technical skills",
        "skills",
        "education",
        r"formaci[oó]n",
        "certifications?",
        "languages?",
        "idiomas",
        "awards?",
        "publications?",
        "volunteer experience",
        "additional information",
    )
    match = re.search(
        rf"(?ims)^\s*({'|'.join(headings)})\s*$([\s\S]*?)(?=^\s*({'|'.join(stop_headings)})\s*$|\Z)",
        text,
    )
    return match.group(2) if match else ""


def _date_range_match(line: str) -> re.Match[str] | None:
    month = (
        r"jan(?:uary)?|feb(?:ruary)?|mar(?:ch)?|apr(?:il)?|may|jun(?:e)?|jul(?:y)?|aug(?:ust)?|"
        r"sep(?:tember)?|sept|oct(?:ober)?|nov(?:ember)?|dec(?:ember)?|"
        r"ene(?:ro)?|feb(?:rero)?|mar(?:zo)?|abr(?:il)?|may(?:o)?|jun(?:io)?|jul(?:io)?|"
        r"ago(?:sto)?|sep(?:tiembre)?|sept(?:iembre)?|oct(?:ubre)?|nov(?:iembre)?|dic(?:iembre)?"
    )
    month_year = rf"(?:{month})\.?\s+\d{{4}}"
    numeric_month_year = r"(?:0?[1-9]|1[0-2])[/.-]\d{4}"
    year_only = r"\d{4}"
    separator = r"\s*[-–—]\s*"
    dated_start = rf"(?:{month_year}|{numeric_month_year}|{year_only})"
    end = rf"(?:{month_year}|{numeric_month_year}|{year_only}|present|current|actualidad|presente)"
    return re.search(rf"(?i)\b{dated_start}{separator}{end}\b", line)


def _next_company_line(lines: list[str], start: int) -> str:
    for line in lines[start : start + 3]:
        stripped = line.strip()
        if not stripped or stripped.startswith(BULLET_PREFIXES):
            continue
        if _date_range_match(stripped):
            continue
        return stripped
    return ""


def _company_match_terms(company: str) -> list[str]:
    normalized = _normalize_for_match(company)
    stopwords = {
        "client",
        "cliente",
        "malaga",
        "spain",
        "espana",
        "buenos",
        "aires",
        "argentina",
        "remote",
        "remoto",
        "consulting",
        "group",
    }
    tokens = [
        token
        for token in re.findall(r"[a-z0-9]+", normalized)
        if len(token) >= 4 and token not in stopwords
    ]
    seen = set()
    unique = []
    for token in tokens:
        if token not in seen:
            seen.add(token)
            unique.append(token)
    return unique[:5]


def _normalize_for_match(text: str) -> str:
    decomposed = unicodedata.normalize("NFKD", str(text or ""))
    ascii_text = "".join(char for char in decomposed if not unicodedata.combining(char))
    return ascii_text.lower()


def _contains_phrase_for_materials(normalized_text: str, phrase: str) -> bool:
    normalized_phrase = _normalize_for_match(phrase)
    if not normalized_phrase:
        return False
    return normalized_phrase in normalized_text


def _to_dict(value: Any) -> dict[str, Any]:
    if isinstance(value, dict):
        return value
    if is_dataclass(value):
        return asdict(value)
    if hasattr(value, "to_dict"):
        return value.to_dict()
    if hasattr(value, "__dict__"):
        return vars(value)
    return {}


def _compact_job(job: dict[str, Any]) -> dict[str, Any]:
    keys = [
        "id",
        "job_id",
        "title",
        "company",
        "location",
        "workplace_type",
        "source",
        "url",
        "apply_url",
        "description_text",
        "posted_at",
        "first_seen_at",
        "last_seen_at",
    ]
    compact = {key: job.get(key) for key in keys if job.get(key) is not None}
    description = str(compact.get("description_text") or "")
    if len(description) > 9000:
        compact["description_text"] = description[:9000] + "\n[truncated]"
    return compact
